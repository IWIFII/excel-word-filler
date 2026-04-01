"""
从 Excel 指定列读取数据，写入 Word 模板中的表格。

为什么很多“填充 Word”效果不好：
1. 在 Word 里用“查找替换”整段替换——占位符可能被拆成多个 run，替换不到或只替换一部分。
2. 直接改 XML——容易破坏段落、编号、域、合并单元格。
3. 合并单元格——python-docx 对合并表支持有限，尽量用规则矩形表格。

推荐做法：用 python-docx 按“单元格”写入，需要多行时复制模板里已有数据行的 XML，
这样字体、边框、对齐会继承那一行。

使用前：
- 将 Excel 放到 source_docs/
- 将 Word 模板放到 templates/，第一列表头行（HEADER_ROW_INDEX）中的文字需与 Excel 表头一致（可多空格/全角括号差异，会自动规范化匹配）。
- 模板表格中保留一行“样例行”（TEMPLATE_DATA_ROW_INDEX），程序会按该行复制出更多行再填值。
- AUTO_MATCH_HEADERS=True（默认）时按 Word 表头顺序从 Excel 取列，无需再手写列名。
"""

from __future__ import annotations

import shutil
import unicodedata
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.table import Table


# =========================================
BASE = Path(__file__).resolve().parent
# 若该路径不存在，会尝试使用 source_docs / templates 下第一个 .xlsx / .docx
EXCEL_PATH = BASE / "source_docs" / "data.xlsx"
WORD_TEMPLATE = BASE / "templates" / "template.docx"
OUTPUT_PATH = BASE / "output" / "filled.docx"

# True：用 Word 表格第 HEADER_ROW_INDEX 行作为列名，与 Excel 第一行表头自动匹配后按该顺序取数。
# False：使用下方 EXCEL_COLUMNS 或 EXCEL_USE_COL_INDEXES。
AUTO_MATCH_HEADERS = True

# 以下为 AUTO_MATCH_HEADERS=False 时使用
EXCEL_COLUMNS: list[str] | None = None
EXCEL_USE_COL_INDEXES: list[int] | None = None

SHEET_NAME: str | None = None  # None 表示第一个 sheet

# Word：第几个表格（0 表示文档里第一个表）
TABLE_INDEX = 0

# 模板表格中：第几行是表头（不复制、不覆盖）；第几行是“样例数据行”（复制此行的格式与结构）
HEADER_ROW_INDEX = 0
TEMPLATE_DATA_ROW_INDEX = 1

# 若 Excel 行数少于模板中表头以下的行，是否删空多余行（谨慎：可能影响合并单元格）
TRIM_EXTRA_TEMPLATE_ROWS = False

# 只读前 N 行（None 表示不按此项限制；实际还会受 MAX_ROWS_TO_FILL 约束）
READ_EXCEL_NROWS: int | None = None

# 写入 Word 的数据最多多少行（None 表示不限制）。大表请设较小值，否则复制表格行极慢。
MAX_ROWS_TO_FILL: int | None = 2000

# 复制 Word 表格行时每多少行打印一次进度（0 表示不打印）
ROW_COPY_PROGRESS_EVERY = 500
# ==========================================


def _effective_read_nrows() -> int | None:
    """传给 pandas 的 nrows：与 MAX_ROWS_TO_FILL 取较小值，避免整表读入内存。"""
    if READ_EXCEL_NROWS is not None and MAX_ROWS_TO_FILL is not None:
        return min(READ_EXCEL_NROWS, MAX_ROWS_TO_FILL)
    if READ_EXCEL_NROWS is not None:
        return READ_EXCEL_NROWS
    return MAX_ROWS_TO_FILL


def _log(msg: str) -> None:
    print(msg, flush=True)


def _first_office_file(folder: Path, patterns: tuple[str, ...]) -> Path | None:
    """返回目录下第一个匹配文件，排除 Excel 打开的临时文件 ~$."""
    if not folder.is_dir():
        return None
    for pattern in patterns:
        for p in sorted(folder.glob(pattern)):
            if p.is_file() and not p.name.startswith("~$"):
                return p
    return None


def _excel_engine(path: Path) -> str:
    """pandas 读表引擎：.xlsx 用 openpyxl，旧版 .xls 用 xlrd。"""
    suf = path.suffix.lower()
    if suf == ".xls":
        return "xlrd"
    return "openpyxl"


def resolve_excel_path() -> Path:
    if EXCEL_PATH.is_file():
        return EXCEL_PATH
    # 顺序：source_docs → 项目根目录（文件常放在这里而不是子文件夹）
    for folder in (BASE / "source_docs", BASE):
        found = _first_office_file(folder, ("*.xlsx", "*.xls"))
        if found is not None:
            return found
    raise FileNotFoundError(
        f"未找到 Excel：默认 {EXCEL_PATH} 不存在；已在 "
        f"{BASE / 'source_docs'} 与 {BASE} 根目录查找 *.xlsx / *.xls 仍无结果。"
        f"请把表格放进上述目录之一，或把 EXCEL_PATH 设为你的文件绝对路径。"
    )


def resolve_word_template_path() -> Path:
    if WORD_TEMPLATE.is_file():
        return WORD_TEMPLATE
    for folder in (BASE / "templates", BASE):
        found = _first_office_file(folder, ("*.docx",))
        if found is not None:
            # 避免误用输出目录里的结果文件
            if found.resolve() == OUTPUT_PATH.resolve():
                continue
            if found.parent.name.lower() == "output":
                continue
            return found
    raise FileNotFoundError(
        f"未找到 Word 模板：默认 {WORD_TEMPLATE} 不存在；已在 "
        f"{BASE / 'templates'} 与 {BASE} 根目录查找 *.docx 仍无结果。"
        f"请把模板放进上述目录之一，或把 WORD_TEMPLATE 设为绝对路径。"
    )


def _normalize_header_text(s: str) -> str:
    """用于表头比对：统一空白、全角括号、Unicode 兼容分解。"""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("（", "(").replace("）", ")")
    s = " ".join(s.split())
    return s


def _word_cell_text(cell) -> str:
    return (cell.text or "").strip()


def _read_word_table_headers(table: Table, header_row_index: int) -> list[str]:
    row = table.rows[header_row_index]
    return [_word_cell_text(c) for c in row.cells]


def _match_word_headers_to_excel(
    word_headers: list[str],
    excel_columns: pd.Index,
) -> list:
    """
    按 Word 表头顺序，返回对应的 Excel 列标签（与 pandas 中一致）。
    先精确匹配整表头字符串，再用规范化后的字符串匹配。
    """
    labels = list(excel_columns)
    if not labels:
        raise ValueError("Excel 没有列名（表头为空）")

    norm_to_label: dict[str, object] = {}
    for lab in labels:
        key = _normalize_header_text(str(lab))
        if key not in norm_to_label:
            norm_to_label[key] = lab

    # 精确匹配：原始 str(lab) 去首尾空白
    strip_to_label: dict[str, object] = {}
    for lab in labels:
        k = str(lab).strip()
        if k not in strip_to_label:
            strip_to_label[k] = lab

    matched: list = []
    for i, wh in enumerate(word_headers):
        if not wh:
            raise ValueError(
                f"Word 表头第 {i + 1} 列为空。请填写与 Excel 一致的列名，或取消合并表头导致的空单元格。"
            )
        if wh in strip_to_label:
            matched.append(strip_to_label[wh])
            continue
        nk = _normalize_header_text(wh)
        if nk in norm_to_label:
            matched.append(norm_to_label[nk])
            continue
        raise ValueError(
            f"Word 表头「{wh}」在 Excel 中找不到对应列。"
            f" Excel 表头为: {[str(x) for x in labels]}"
        )
    return matched


def _duplicate_table_row(table: Table, source_row_index: int) -> None:
    """在 source_row_index 下方插入一行，结构与样式与 source_row_index 相同。"""
    tbl = table._tbl
    tr = table.rows[source_row_index]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)


def fill_word_table_from_excel() -> Path:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    excel_file = resolve_excel_path()
    word_file = resolve_word_template_path()
    if excel_file != EXCEL_PATH:
        _log(f"使用 Excel: {excel_file}")
    if word_file != WORD_TEMPLATE:
        _log(f"使用模板: {word_file}")

    eng = _excel_engine(excel_file)
    read_kw: dict = {
        "sheet_name": SHEET_NAME or 0,
        "engine": eng,
    }
    _nrows = _effective_read_nrows()
    if _nrows is not None:
        read_kw["nrows"] = _nrows

    _log(
        "正在读取 Excel（文件很大或正用 Excel 打开时可能很慢；可先关闭 Excel 再试）..."
    )
    try:
        if AUTO_MATCH_HEADERS:
            df_full = pd.read_excel(excel_file, header=0, **read_kw)
            _log("正在读取 Word 模板表头并与 Excel 表头匹配...")
            doc_preview = Document(word_file)
            if TABLE_INDEX >= len(doc_preview.tables):
                raise IndexError(
                    f"文档只有 {len(doc_preview.tables)} 个表格，TABLE_INDEX={TABLE_INDEX} 无效"
                )
            word_headers = _read_word_table_headers(
                doc_preview.tables[TABLE_INDEX], HEADER_ROW_INDEX
            )
            matched_labels = _match_word_headers_to_excel(word_headers, df_full.columns)
            df = df_full.loc[:, matched_labels]
            col_seq = list(df.columns)
            _log("表头匹配（Word → Excel 列）：")
            for wh, lab in zip(word_headers, matched_labels):
                _log(f"  「{wh}」 → 「{lab}」")
        elif EXCEL_COLUMNS:
            df = pd.read_excel(
                excel_file,
                usecols=EXCEL_COLUMNS,
                **read_kw,
            )
            col_seq = list(df.columns)
        elif EXCEL_USE_COL_INDEXES is not None:
            df = pd.read_excel(excel_file, header=0, **read_kw)
            df = df.iloc[:, EXCEL_USE_COL_INDEXES]
            col_seq = list(df.columns)
        else:
            raise ValueError(
                "请设置 AUTO_MATCH_HEADERS=True，或设置 EXCEL_COLUMNS / EXCEL_USE_COL_INDEXES"
            )
    except ValueError as e:
        if AUTO_MATCH_HEADERS:
            raise
        try:
            peek = pd.read_excel(
                excel_file,
                sheet_name=SHEET_NAME or 0,
                engine=eng,
                header=0,
                nrows=0,
            )
            heads = list(peek.columns)
        except Exception:
            heads = "(无法读取表头)"
        raise ValueError(
            "读取列失败，请检查 EXCEL_COLUMNS 是否与表头完全一致。"
            f" 当前文件表头为: {heads}。原始错误: {e}"
        ) from e

    if MAX_ROWS_TO_FILL is not None:
        df = df.head(MAX_ROWS_TO_FILL)
        _log(
            f"已按 MAX_ROWS_TO_FILL={MAX_ROWS_TO_FILL} 限制为 {len(df)} 行数据（不超过该上限）。"
        )

    rows_data = df.astype(object).where(pd.notna(df), "").astype(str).values.tolist()
    n = len(rows_data)
    if n == 0:
        raise ValueError("Excel 没有数据行")

    _log(f"已读取 Excel：{n} 行。正在打开 Word 模板...")
    shutil.copy2(word_file, OUTPUT_PATH)
    doc = Document(OUTPUT_PATH)
    if TABLE_INDEX >= len(doc.tables):
        raise IndexError(f"文档只有 {len(doc.tables)} 个表格，TABLE_INDEX={TABLE_INDEX} 无效")

    table = doc.tables[TABLE_INDEX]
    n_cols_word = len(table.rows[TEMPLATE_DATA_ROW_INDEX].cells)
    n_cols_excel = len(col_seq)
    if n_cols_excel > n_cols_word:
        raise ValueError(
            f"Excel 列数 ({n_cols_excel}) 多于 Word 表该行单元格数 ({n_cols_word})，请调整模板列数或选列"
        )

    if n > 2000 and MAX_ROWS_TO_FILL is None:
        _log(
            f"提示：共 {n} 行，复制 Word 表格行可能需数分钟；"
            f"可设置 MAX_ROWS_TO_FILL 或 READ_EXCEL_NROWS 限制行数。"
        )

    # 需要的数据行数 = n；模板里已有 1 行样例，每次复制「当前最后一行」以保证顺序正确
    last_data_row = TEMPLATE_DATA_ROW_INDEX
    total_copies = n - 1
    for i in range(total_copies):
        _duplicate_table_row(table, last_data_row)
        last_data_row += 1
        if ROW_COPY_PROGRESS_EVERY and (i + 1) % ROW_COPY_PROGRESS_EVERY == 0:
            _log(f"  已复制表格行 {i + 1} / {total_copies} ...")

    # 若希望删掉模板里预置的多余空行（在复制之后索引会变，一般不建议开）
    if TRIM_EXTRA_TEMPLATE_ROWS and n < len(table.rows) - HEADER_ROW_INDEX:
        # 保留表头 + n 行数据，删除下面多余的行（实现略复杂，默认关闭）
        pass

    for r, row_values in enumerate(rows_data):
        word_row_index = TEMPLATE_DATA_ROW_INDEX + r
        row_cells = table.rows[word_row_index].cells
        for c, val in enumerate(row_values):
            if c >= len(row_cells):
                break
            # 整格替换为纯文本；行级样式来自复制的那一行，单元格内若需混排格式需再改用 runs API
            row_cells[c].text = str(val)

    _log("正在保存 Word ...")
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = fill_word_table_from_excel()
    _log(f"已生成: {out}")
